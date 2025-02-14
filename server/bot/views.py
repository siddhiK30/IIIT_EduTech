# # Django Backend (views.py)
# from rest_framework.views import APIView
# from rest_framework.response import Response
# from rest_framework import status
# from django.core.files.storage import default_storage
# import PyPDF2
# import docx
# import os

# class FileProcessView(APIView):
#     def post(self, request):
#         try:
#             file = request.FILES.get('file')
#             if not file:
#                 return Response({'error': 'No file provided'}, status=status.HTTP_400_BAD_REQUEST)

#             # Get file extension
#             file_extension = os.path.splitext(file.name)[1].lower()
#             extracted_text = ""

#             # Process different file types
#             if file_extension == '.pdf':
#                 pdf_reader = PyPDF2.PdfReader(file)
#                 for page in pdf_reader.pages:
#                     extracted_text += page.extract_text()
            
#             elif file_extension == '.docx':
#                 doc = docx.Document(file)
#                 for para in doc.paragraphs:
#                     extracted_text += para.text + "\n"
            
#             elif file_extension == '.txt':
#                 extracted_text = file.read().decode('utf-8')
            
#             else:
#                 return Response(
#                     {'error': 'Unsupported file format'},
#                     status=status.HTTP_400_BAD_REQUEST
#                 )

#             # Clean the extracted text
#             cleaned_text = self.clean_text(extracted_text)
            
#             return Response({
#                 'text': cleaned_text,
#                 'filename': file.name
#             }, status=status.HTTP_200_OK)

#         except Exception as e:
#             return Response(
#                 {'error': str(e)},
#                 status=status.HTTP_500_INTERNAL_SERVER_ERROR
#             )

#     def clean_text(self, text):
#         # Add your text cleaning logic here
#         cleaned = text.strip()
#         cleaned = ' '.join(cleaned.split())
#         return cleaned



from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from django.core.files.storage import default_storage
import PyPDF2
import docx
import os
import google.generativeai as genai

# Configure Gemini API
genai.configure(api_key="AIzaSyApZd8uGx6WJ0EGHEf9xelRoMaSojFrycg")

class FileProcessView(APIView):
    def post(self, request):
        try:
            file = request.FILES.get('file')
            if not file:
                return Response({'error': 'No file provided'}, status=status.HTTP_400_BAD_REQUEST)

            # Get file extension
            file_extension = os.path.splitext(file.name)[1].lower()
            extracted_text = ""

            # Process different file types
            if file_extension == '.pdf':
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    extracted_text += page.extract_text()
            
            elif file_extension == '.docx':
                doc = docx.Document(file)
                for para in doc.paragraphs:
                    extracted_text += para.text + "\n"
            
            elif file_extension == '.txt':
                extracted_text = file.read().decode('utf-8')
            
            else:
                return Response(
                    {'error': 'Unsupported file format'},
                    status=status.HTTP_400_BAD_REQUEST
                )

            # Clean the extracted text
            cleaned_text = self.clean_text(extracted_text)
            
            # Store text in session for context
            request.session['document_text'] = cleaned_text

            # Generate initial summary using Gemini
            model = genai.GenerativeModel('gemini-pro')
            prompt = f"Please analyze and provide concise answers related to the question below"
            response = model.generate_content(prompt)
            
            return Response({
                'text': cleaned_text,
                'summary': response.text,
                'filename': file.name
            }, status=status.HTTP_200_OK)

        except Exception as e:
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def clean_text(self, text):
        # Add your text cleaning logic here
        cleaned = text.strip()
        cleaned = ' '.join(cleaned.split())
        return cleaned

class QuestionAnswerView(APIView):
    def post(self, request):
        try:
            question = request.data.get('question')
            document_text = request.session.get('document_text')

            if not question:
                return Response(
                    {'error': 'No question provided'},
                    status=status.HTTP_400_BAD_REQUEST
                )

            if not document_text:
                # If no context, just use Gemini directly
                model = genai.GenerativeModel('gemini-pro')
                response = model.generate_content(question)
            else:
                # Use document context for answering
                model = genai.GenerativeModel('gemini-pro')
                prompt = f"""
                Based on the following document content:
                {document_text}

                Please answer this question:
                {question}
                """
                response = model.generate_content(prompt)

            return Response({
                'answer': response.text,
                'status': 'success'
            })

        except Exception as e:
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )